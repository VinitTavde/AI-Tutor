import os
import tempfile
from typing import Optional, List, Dict, Tuple
import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:    
    def __init__(self):
        # Keep existing supported formats
        self.text_formats = ['.pdf', '.docx', '.txt']
        
        # Initialize audio processor
        try:
            from audio_processor import AudioProcessor
            self.audio_processor = AudioProcessor()
            self.audio_formats = list(self.audio_processor.supported_formats)
            logger.info("Audio processing enabled")
        except ImportError as e:
            self.audio_processor = None
            self.audio_formats = []
            logger.warning(f"Audio processing disabled: {e}")
        
        # Combined supported formats
        self.supported_formats = self.text_formats + self.audio_formats
    
    def is_audio_file(self, filename: str) -> bool:
        """Check if file is an audio format."""
        if not self.audio_processor:
            return False
        return self.audio_processor.is_audio_file(filename)
    
    def is_file_format(self, filename: str) -> bool:
        """Check if file is a document format (PDF, DOCX, TXT) that supports page extraction."""
        file_extension = os.path.splitext(filename)[1].lower()
        return file_extension in self.text_formats
    
    def extract_pages_from_pdf(self, file_path: str) -> List[Dict]:
        """Extract text from PDF page by page."""
        try:
            pages_data = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text().strip()
                    if page_text:  # Only add non-empty pages
                        pages_data.append({
                            'page_number': page_num,
                            'text': page_text,
                            'total_pages': total_pages
                        })
                
                logger.info(f"Extracted {len(pages_data)} non-empty pages from PDF")
                return pages_data
                
        except Exception as e:
            logger.error(f"Error extracting pages from PDF: {str(e)}")
            raise ValueError(f"Failed to extract pages from PDF: {str(e)}")
    
    def extract_pages_from_docx(self, file_path: str) -> List[Dict]:
        """Extract text from DOCX page by page (approximated by sections/paragraphs)."""
        try:
            doc = Document(file_path)
            pages_data = []
            current_page = 1
            current_page_text = ""
            paragraph_count = 0
            paragraphs_per_page = 20  # Approximate paragraphs per page
            
            total_paragraphs = len(doc.paragraphs)
            estimated_pages = max(1, (total_paragraphs + paragraphs_per_page - 1) // paragraphs_per_page)
            
            for paragraph in doc.paragraphs:
                current_page_text += paragraph.text + "\n"
                paragraph_count += 1
                
                # Create a new "page" every N paragraphs or at the end
                if paragraph_count >= paragraphs_per_page or paragraph == doc.paragraphs[-1]:
                    if current_page_text.strip():
                        pages_data.append({
                            'page_number': current_page,
                            'text': current_page_text.strip(),
                            'total_pages': estimated_pages
                        })
                    
                    current_page += 1
                    current_page_text = ""
                    paragraph_count = 0
            
            # Update total pages with actual count
            for page_data in pages_data:
                page_data['total_pages'] = len(pages_data)
            
            logger.info(f"Extracted {len(pages_data)} pages from DOCX")
            return pages_data
            
        except Exception as e:
            logger.error(f"Error extracting pages from DOCX: {str(e)}")
            raise ValueError(f"Failed to extract pages from DOCX: {str(e)}")
    
    def extract_pages_from_txt(self, file_path: str) -> List[Dict]:
        """Extract text from TXT file by creating artificial pages."""
        try:
            # Read the file
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    full_text = file.read().strip()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as file:
                    full_text = file.read().strip()
            
            if not full_text:
                return []
            
            # Split into artificial pages (e.g., every 2000 characters)
            chars_per_page = 2000
            pages_data = []
            page_num = 1
            
            for i in range(0, len(full_text), chars_per_page):
                page_text = full_text[i:i + chars_per_page]
                if page_text.strip():
                    pages_data.append({
                        'page_number': page_num,
                        'text': page_text.strip(),
                        'total_pages': 0  # Will be updated below
                    })
                    page_num += 1
            
            # Update total pages
            total_pages = len(pages_data)
            for page_data in pages_data:
                page_data['total_pages'] = total_pages
            
            logger.info(f"Extracted {len(pages_data)} artificial pages from TXT")
            return pages_data
            
        except Exception as e:
            logger.error(f"Error extracting pages from TXT: {str(e)}")
            raise ValueError(f"Failed to extract pages from TXT: {str(e)}")

    def extract_text_from_pdf(self, file_path: str) -> str:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Error extracting text from TXT: {str(e)}")
                raise ValueError(f"Failed to extract text from TXT: {str(e)}")
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
    
    def extract_text_from_audio(self, file_path: str, language: Optional[str] = None) -> str:
        """Extract text from audio file via transcription."""
        if not self.audio_processor:
            raise ValueError("Audio processing not available. Please install required dependencies.")
        
        try:
            transcribed_text, metadata = self.audio_processor.transcribe_audio(file_path, language)
            
            # Log transcription metadata for debugging
            logger.info(f"Audio transcription metadata: {metadata}")
            
            # Store metadata for potential later use (could be passed back if needed)
            self._last_audio_metadata = metadata
            
            return transcribed_text
            
        except Exception as e:
            logger.error(f"Error transcribing audio: {str(e)}")
            raise ValueError(f"Failed to transcribe audio: {str(e)}")
    
    def process_document_with_pages(self, file_path: str, file_name: str, chunk_size: int = 1000, overlap: int = 200) -> Tuple[List[Dict], int]:
        """
        Process document page by page and create chunks with page information.
        Only works for file formats (PDF, DOCX, TXT).
        
        Args:
            file_path: Path to the file
            file_name: Name of the file
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            Tuple of (list of chunk dictionaries with page info, total_pages)
        """
        file_extension = os.path.splitext(file_name)[1].lower()
        
        if file_extension not in self.text_formats:
            raise ValueError(f"Page-based processing only supported for: {', '.join(self.text_formats)}")
        
        # Extract pages based on file type
        if file_extension == '.pdf':
            pages_data = self.extract_pages_from_pdf(file_path)
        elif file_extension == '.docx':
            pages_data = self.extract_pages_from_docx(file_path)
        elif file_extension == '.txt':
            pages_data = self.extract_pages_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format for page processing: {file_extension}")
        
        if not pages_data:
            raise ValueError("No readable pages found in document")
        
        # Process each page and create chunks
        all_chunks = []
        chunk_index = 0
        total_pages = pages_data[0]['total_pages'] if pages_data else 0
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=len,
            is_separator_regex=False,
        )
        
        for page_data in pages_data:
            page_text = page_data['text']
            page_number = page_data['page_number']
            
            # Split page text into chunks
            page_chunks = text_splitter.split_text(page_text)
            
            # Create chunk objects with page information
            for chunk_text in page_chunks:
                chunk_dict = {
                    'text': chunk_text,
                    'page_number': page_number,
                    'chunk_index': chunk_index,
                    'total_pages': total_pages
                }
                all_chunks.append(chunk_dict)
                chunk_index += 1
        
        logger.info(f"Created {len(all_chunks)} chunks from {total_pages} pages")
        return all_chunks, total_pages

    def process_document(self, file_path: str, file_name: str, language: Optional[str] = None) -> str:
        """
        Process document (text or audio) and extract text content.
        This is the legacy method for non-page-based processing.
        
        Args:
            file_path: Path to the file
            file_name: Name of the file
            language: Language code for audio transcription (optional)
            
        Returns:
            Extracted text content
        """
        file_extension = os.path.splitext(file_name)[1].lower()
        
        if file_extension not in self.supported_formats:
            supported_text = ', '.join(self.text_formats)
            supported_audio = ', '.join(self.audio_formats) if self.audio_formats else 'None (audio processing not available)'
            raise ValueError(f"Unsupported file format: {file_extension}. "
                           f"Supported text formats: {supported_text}. "
                           f"Supported audio formats: {supported_audio}")
        
        # Handle audio files
        if self.is_audio_file(file_name):
            return self.extract_text_from_audio(file_path, language)
        
        # Handle text files (existing logic unchanged)
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def get_last_audio_metadata(self) -> Optional[Dict]:
        """Get metadata from the last audio transcription."""
        return getattr(self, '_last_audio_metadata', None)
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for better embedding."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=len,
            is_separator_regex=False,
        )
        return text_splitter.split_text(text)
    